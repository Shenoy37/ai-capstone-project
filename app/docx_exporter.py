from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, Iterable

from docx import Document


def _add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item)
        paragraph.style = "List Bullet"


def build_brd_document(brd_payload: Dict[str, Any]) -> BytesIO:
    """Create an in-memory Word document from the BRD payload."""

    document = Document()
    document.add_heading(brd_payload.get("title", "Business Requirements Document"), level=0)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(brd_payload.get("executive_summary", ""))

    document.add_heading("Objectives", level=1)
    _add_bullet_list(document, brd_payload.get("objectives", []))

    scope = brd_payload.get("scope", {})
    document.add_heading("Scope", level=1)
    document.add_heading("In Scope", level=2)
    _add_bullet_list(document, scope.get("in_scope", []))
    document.add_heading("Out of Scope", level=2)
    _add_bullet_list(document, scope.get("out_of_scope", []))

    document.add_heading("Stakeholders", level=1)
    _add_bullet_list(document, brd_payload.get("stakeholders", []))

    document.add_heading("Functional Requirements", level=1)
    _add_bullet_list(document, brd_payload.get("functional_requirements", []))

    document.add_heading("Non-Functional Requirements", level=1)
    _add_bullet_list(document, brd_payload.get("non_functional_requirements", []))

    document.add_heading("Data Requirements", level=1)
    _add_bullet_list(document, brd_payload.get("data_requirements", []))

    document.add_heading("Regulatory & Compliance", level=1)
    _add_bullet_list(document, brd_payload.get("regulatory_and_compliance", []))

    document.add_heading("Risks & Mitigations", level=1)
    _add_bullet_list(document, brd_payload.get("risks_and_mitigations", []))

    document.add_heading("Timeline", level=1)
    _add_bullet_list(document, brd_payload.get("timeline", []))

    document.add_heading("Dependencies", level=1)
    _add_bullet_list(document, brd_payload.get("dependencies", []))

    document.add_heading("Open Questions", level=1)
    _add_bullet_list(document, brd_payload.get("open_questions", []))

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
