"""
Document Service for BRD Generator - Handles document generation and export
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List
import os
from datetime import datetime
import logging
from config.settings import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentService:
    """Service class for document generation and export"""
    
    def __init__(self):
        """Initialize the document service"""
        self.output_dir = Config.GENERATED_DIR
        os.makedirs(self.output_dir, exist_ok=True)
        logger.info("Document Service initialized successfully")
    
    def create_brd_document(self, 
                           brd_sections: Dict[str, str],
                           domain: str,
                           project_title: str,
                           metadata: Dict = None) -> str:
        """
        Create a BRD document in Word format
        
        Args:
            brd_sections: Dictionary containing BRD sections
            domain: Business domain
            project_title: Title of the project
            metadata: Additional metadata for the document
            
        Returns:
            File path of the generated document
        """
        
        try:
            # Create new document
            doc = Document()
            
            # Add title page
            self._add_title_page(doc, project_title, domain, metadata)
            
            # Add table of contents
            self._add_table_of_contents(doc, list(brd_sections.keys()))
            
            # Add BRD sections
            self._add_brd_sections(doc, brd_sections)
            
            # Add appendix
            self._add_appendix(doc, domain)
            
            # Generate filename and save
            filename = self._generate_filename(project_title, domain)
            filepath = os.path.join(self.output_dir, filename)
            
            doc.save(filepath)
            logger.info(f"BRD document created successfully: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error creating BRD document: {str(e)}")
            raise
    
    def _add_title_page(self, doc: Document, project_title: str, domain: str, metadata: Dict = None):
        """Add title page to the document"""
        
        # Add main title
        title = doc.add_heading('Business Requirement Document', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add project title
        project_heading = doc.add_heading(project_title, 1)
        project_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add domain
        domain_para = doc.add_paragraph()
        domain_para.add_run(f"Domain: {domain}").bold = True
        domain_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        if metadata:
            doc.add_paragraph()  # Add space
            
            for key, value in metadata.items():
                meta_para = doc.add_paragraph()
                meta_para.add_run(f"{key}: ").bold = True
                meta_para.add_run(str(value))
                meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add generation date
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document, sections: List[str]):
        """Add table of contents"""
        
        # Add TOC title
        toc_heading = doc.add_heading('Table of Contents', 1)
        
        # Add section list
        for i, section in enumerate(sections, 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. ").bold = True
            para.add_run(section)
            para.style = 'List Number'
        
        # Add page break
        doc.add_page_break()
    
    def _add_brd_sections(self, doc: Document, brd_sections: Dict[str, str]):
        """Add all BRD sections to the document"""
        
        section_number = 1
        
        for section_name, section_content in brd_sections.items():
            # Add section heading
            heading = doc.add_heading(f"{section_number}. {section_name}", 1)
            
            # Add section content
            if section_content:
                # Split content into paragraphs
                paragraphs = section_content.split('\n')
                
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Check if it's a bullet point or numbered list
                        if para_text.startswith('-') or para_text.startswith('*'):
                            para = doc.add_paragraph(para_text[1:].strip(), style='List Bullet')
                        elif para_text[0].isdigit() and '.' in para_text[:3]:
                            para = doc.add_paragraph(para_text, style='List Number')
                        else:
                            para = doc.add_paragraph(para_text)
            else:
                doc.add_paragraph("[Content to be developed]")
            
            # Add some space between sections
            doc.add_paragraph()
            section_number += 1
    
    def _add_appendix(self, doc: Document, domain: str):
        """Add appendix with domain-specific information"""
        
        # Add appendix heading
        doc.add_page_break()
        appendix_heading = doc.add_heading('Appendix', 1)
        
        # Add compliance references
        compliance_heading = doc.add_heading('Compliance References', 2)
        
        if domain.lower() == "pharma":
            compliance_items = [
                "FDA 21 CFR Part 11 - Electronic Records and Signatures",
                "HIPAA - Health Insurance Portability and Accountability Act",
                "GxP - Good Practice Guidelines",
                "Clinical Trial Regulations (21 CFR Part 312)",
                "Adverse Event Reporting Requirements (21 CFR Part 314.80)"
            ]
        elif domain.lower() == "finance":
            compliance_items = [
                "Basel III - International Regulatory Framework for Banks",
                "GDPR - General Data Protection Regulation",
                "SOX - Sarbanes-Oxley Act",
                "Risk Management Guidelines",
                "Data Privacy and Security Standards"
            ]
        else:
            compliance_items = ["Standard compliance requirements"]
        
        for item in compliance_items:
            para = doc.add_paragraph(item, style='List Bullet')
        
        # Add glossary
        doc.add_paragraph()
        glossary_heading = doc.add_heading('Glossary', 2)
        
        glossary_text = """
        This document contains domain-specific terminology relevant to the {domain} industry.
        For detailed definitions and explanations, please refer to the respective regulatory
        guidelines and industry standards mentioned in the compliance references.
        """.format(domain=domain)
        
        doc.add_paragraph(glossary_text)
    
    def _generate_filename(self, project_title: str, domain: str) -> str:
        """Generate a unique filename for the BRD document"""
        
        # Clean project title for filename
        clean_title = "".join(c for c in project_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_title = clean_title.replace(' ', '_')
        
        # Add timestamp and domain
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"BRD_{domain}_{clean_title}_{timestamp}.docx"
        
        return filename
    
    def export_to_pdf(self, docx_path: str) -> str:
        """
        Export DOCX to PDF (placeholder for future implementation)
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Path to the generated PDF file
        """
        
        # This is a placeholder for PDF export functionality
        # In a real implementation, you would use libraries like reportlab or comtypes
        logger.info("PDF export not yet implemented")
        return docx_path
    
    def get_document_preview(self, brd_sections: Dict[str, str]) -> str:
        """
        Generate a text preview of the BRD document
        
        Args:
            brd_sections: Dictionary containing BRD sections
            
        Returns:
            Formatted text preview
        """
        
        preview_lines = []
        preview_lines.append("=== BRD DOCUMENT PREVIEW ===\n")
        
        for section_name, section_content in brd_sections.items():
            preview_lines.append(f"## {section_name}")
            
            if section_content:
                # Show first 200 characters of each section
                preview = section_content[:200] + "..." if len(section_content) > 200 else section_content
                preview_lines.append(preview)
            else:
                preview_lines.append("[Content to be developed]")
            
            preview_lines.append("\n" + "-" * 50 + "\n")
        
        return "\n".join(preview_lines)